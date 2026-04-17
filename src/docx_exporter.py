import logging
from typing import cast
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.styles.style import _ParagraphStyle

logger = logging.getLogger(__name__)

# Professional color scheme
BRAND_BLUE = RGBColor(0, 32, 96)      # Dark blue
ACCENT_BLUE = RGBColor(31, 78, 121)   # Medium blue
LIGHT_GRAY = RGBColor(242, 242, 242)  # Background
TEXT_DARK = RGBColor(45, 45, 45)      # Dark gray

class SOCDocxExporter:
    def __init__(self, output_path: Path):
        self.output_path = output_path
        self.doc = Document()
        self._setup_styles()
        self._set_margins()

    def _set_margins(self):
        """Set professional document margins."""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

    def _shading_elm(self, shd_color):
        """Helper to add cell shading."""
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), shd_color)
        return shd

    def _setup_styles(self):
        """Enhanced corporate styling with professional hierarchy."""
        
        # Normal text
        if 'Normal' in self.doc.styles:
            style = cast(_ParagraphStyle, self.doc.styles['Normal'])
            style.font.name = 'Calibri'
            style.font.size = Pt(11)
            style.font.color.rgb = TEXT_DARK
            style.paragraph_format.line_spacing = 1.15
            style.paragraph_format.space_after = Pt(6)

        # Title
        if 'Title' in self.doc.styles:
            title_style = cast(_ParagraphStyle, self.doc.styles['Title'])
            title_style.font.name = 'Calibri'
            title_style.font.size = Pt(28)
            title_style.font.bold = True
            title_style.font.color.rgb = BRAND_BLUE
            title_style.paragraph_format.space_after = Pt(12)

        # Heading 1 (Major sections)
        if 'Heading 1' in self.doc.styles:
            h1 = cast(_ParagraphStyle, self.doc.styles['Heading 1'])
            h1.font.name = 'Calibri'
            h1.font.size = Pt(16)
            h1.font.bold = True
            h1.font.color.rgb = BRAND_BLUE
            h1.paragraph_format.space_before = Pt(12)
            h1.paragraph_format.space_after = Pt(8)

        # Heading 2 (Subsections)
        if 'Heading 2' in self.doc.styles:
            h2 = cast(_ParagraphStyle, self.doc.styles['Heading 2'])
            h2.font.name = 'Calibri'
            h2.font.size = Pt(13)
            h2.font.bold = True
            h2.font.color.rgb = ACCENT_BLUE
            h2.paragraph_format.space_before = Pt(10)
            h2.paragraph_format.space_after = Pt(6)

    def _add_footer(self):
        """Adds professional footer with page numbers and confidentiality notice."""
        section = self.doc.sections[0]
        footer = section.footer
        
        # Clear default and create two-line footer
        footer.paragraphs[0].text = ""
        
        # Add horizontal line
        p_line = footer.paragraphs[0]
        pPr = p_line._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '001F4F')
        pBdr.append(bottom)
        pPr.append(pBdr)
        
        # Confidentiality text
        p_conf = footer.add_paragraph()
        p_conf.text = "Confidential Document - For Internal Use Only"
        p_conf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_conf.runs[0].font.size = Pt(8)
        p_conf.runs[0].font.italic = True
        p_conf.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        
        # Add page number on right
        p_page = footer.add_paragraph()
        p_page.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p_page.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)
        
        run = p_page.add_run()
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE \\* MERGEFORMAT '
        run._r.append(instrText)
        
        run = p_page.add_run()
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar2)
        
        p_page.runs[0].font.size = Pt(8)
        p_page.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    def _add_cover_page(self, run_id: str, hostname: str):
        """Generates a professional cover page with visual hierarchy."""
        
        # Logo/Brand area
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('iSecurify')
        run.font.size = Pt(32)
        run.font.bold = True
        run.font.color.rgb = BRAND_BLUE
        title.paragraph_format.space_after = Pt(2)
        
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run('Cybersecurity Forensic Investigation Report')
        run.font.size = Pt(14)
        run.font.color.rgb = ACCENT_BLUE
        subtitle.paragraph_format.space_after = Pt(24)
        
        # Spacer
        for _ in range(2):
            self.doc.add_paragraph()
        
        # Report metadata in professional format
        metadata_table = self.doc.add_table(rows=5, cols=2)
        metadata_table.autofit = False
        metadata_table.allow_autofit = False
        
        # Set column widths
        metadata_table.columns[0].width = Inches(1.5)
        metadata_table.columns[1].width = Inches(4.0)
        
        metadata = [
            ("Investigation Target:", hostname),
            ("Investigation ID:", run_id),
            ("Document Type:", "Executive Investigation Report"),
            ("Classification:", "Confidential"),
            ("Generated:", "iSecurify Automated SOC Agent")
        ]
        
        for i, (label, value) in enumerate(metadata):
            row = metadata_table.rows[i]
            
            # Left cell (label) - blue background
            left_cell = row.cells[0]
            left_cell._element.get_or_add_tcPr().append(self._shading_elm('001F4F'))
            p_label = left_cell.paragraphs[0]
            p_label.text = label
            p_label.runs[0].font.bold = True
            p_label.runs[0].font.size = Pt(10)
            p_label.runs[0].font.color.rgb = RGBColor(255, 255, 255)
            p_label.paragraph_format.space_before = Pt(6)
            p_label.paragraph_format.space_after = Pt(6)
            
            # Right cell (value) - light background
            right_cell = row.cells[1]
            right_cell._element.get_or_add_tcPr().append(self._shading_elm('F2F2F2'))
            p_value = right_cell.paragraphs[0]
            p_value.text = value
            p_value.runs[0].font.size = Pt(11)
            p_value.paragraph_format.space_before = Pt(6)
            p_value.paragraph_format.space_after = Pt(6)
        
        # Footer on cover page
        self.doc.add_paragraph()
        footer_p = self.doc.add_paragraph()
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_p.add_run('━' * 60)
        run.font.color.rgb = BRAND_BLUE
        run.font.size = Pt(8)
        
        final_p = self.doc.add_paragraph()
        final_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = final_p.add_run('This document contains confidential and sensitive information.\nUnauthorized distribution is strictly prohibited.')
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)
        
        self.doc.add_page_break()

    def _add_section_break(self, title: str):
        """Add a professional section break with colored bar."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(0)
        
        # Colored bar
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '24')
        top.set(qn('w:space'), '0')
        top.set(qn('w:color'), '001F4F')
        pBdr.append(top)
        pPr.append(pBdr)
        
        # Title
        p_title = self.doc.add_heading(title, level=1)
        p_title.paragraph_format.space_before = Pt(6)
        p_title.paragraph_format.space_after = Pt(12)

    def _format_markdown_text(self, text: str):
        """Parse markdown formatting and return formatted paragraph."""
        p = self.doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(8)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Simple markdown parsing
        parts = []
        current = ""
        i = 0
        while i < len(text):
            # Bold **text**
            if text[i:i+2] == '**':
                if current:
                    parts.append(('normal', current))
                    current = ""
                i += 2
                bold_text = ""
                while i < len(text) and text[i:i+2] != '**':
                    bold_text += text[i]
                    i += 1
                parts.append(('bold', bold_text))
                i += 2
            # Italic *text*
            elif text[i] == '*' and (not current or current[-1] != '*'):
                if current:
                    parts.append(('normal', current))
                    current = ""
                i += 1
                italic_text = ""
                while i < len(text) and text[i] != '*':
                    italic_text += text[i]
                    i += 1
                parts.append(('italic', italic_text))
                i += 1
            else:
                current += text[i]
                i += 1
        
        if current:
            parts.append(('normal', current))
        
        # Apply formatting
        for style, content in parts:
            if content.strip():
                run = p.add_run(content)
                run.font.size = Pt(11)
                run.font.color.rgb = TEXT_DARK
                if style == 'bold':
                    run.bold = True
                elif style == 'italic':
                    run.italic = True
        
        return p

    def generate(self, report_text: str, run_id: str, hostname: str = "Production Server"):
        """Converts phased Markdown into a professional Word document."""
        self._add_cover_page(run_id, hostname)
        self._add_footer()

        current_level = 0
        
        for block in report_text.split('\n\n'):
            block = block.strip()
            if not block:
                continue

            # Section Headers with hierarchy
            if block.startswith('#'):
                heading_level = len(block) - len(block.lstrip('#'))
                clean_text = block.replace('#', '').strip()
                
                if heading_level == 1:
                    self._add_section_break(clean_text)
                else:
                    # Sub-heading
                    h = self.doc.add_heading(clean_text, level=min(heading_level, 3))
                    h.paragraph_format.space_before = Pt(8)
                    h.paragraph_format.space_after = Pt(6)
            
            # Bullet lists with better formatting
            elif block.startswith('-'):
                for line in block.split('\n'):
                    line = line.strip()
                    if line.startswith('-'):
                        line = line[1:].strip()
                    
                    if line:
                        # Parse for key: value pairs
                        if ':' in line:
                            key, val = line.split(':', 1)
                            p = self.doc.add_paragraph(style='List Bullet')
                            p.paragraph_format.left_indent = Inches(0.5)
                            p.paragraph_format.space_after = Pt(4)
                            
                            run_key = p.add_run(f"{key.strip()}: ")
                            run_key.bold = True
                            run_key.font.color.rgb = ACCENT_BLUE
                            
                            run_val = p.add_run(val.strip())
                            run_val.font.color.rgb = TEXT_DARK
                        else:
                            p = self.doc.add_paragraph(line, style='List Bullet')
                            p.paragraph_format.left_indent = Inches(0.5)
                            p.paragraph_format.space_after = Pt(4)
            
            # Tables (simple detection)
            elif '|' in block and block.count('|') > 2:
                rows = [row.strip() for row in block.split('\n') if '|' in row]
                if len(rows) > 1:
                    # Parse table
                    headers = [cell.strip() for cell in rows[0].split('|') if cell.strip()]
                    table = self.doc.add_table(rows=len(rows), cols=len(headers))
                    table.style = 'Light Grid Accent 1'
                    
                    # Header row
                    for col_idx, header in enumerate(headers):
                        cell = table.rows[0].cells[col_idx]
                        cell._element.get_or_add_tcPr().append(self._shading_elm('001F4F'))
                        p = cell.paragraphs[0]
                        p.text = header
                        p.runs[0].font.bold = True
                        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
                    
                    # Data rows
                    for row_idx, row in enumerate(rows[1:], 1):
                        cells_data = [cell.strip() for cell in row.split('|') if cell.strip()]
                        for col_idx, cell_data in enumerate(cells_data):
                            if col_idx < len(table.rows[row_idx].cells):
                                table.rows[row_idx].cells[col_idx].text = cell_data
                    
                    self.doc.add_paragraph()  # Space after table
                else:
                    self._format_markdown_text(block)
            
            # Regular paragraphs
            else:
                self._format_markdown_text(block)

        self.doc.save(str(self.output_path))
        logger.info(f"✓ Professional iSecurify Report Saved: {self.output_path.name}")